import re
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
import pytz
import pandas as pd
import streamlit as st


st.set_page_config(page_title="REDCap Formatter", layout="wide")
st.title("🔄 REDCap Instruments Formatter")

# choose which instrument you want to format
instrument = st.sidebar.selectbox(
    "Select instrument", 
    ["OASIS Evaluation", "Checklist Entry", "Email Record Mapper", "NBME Scores", "Preceptor Matching", "Roster_HMC", "Roster_KP", "SDOH Form", "Developmental Assessment Form", "Weekly Quiz Reports", "Documentation Submission #1", "Documentation Submission #2"]
)

if instrument == "OASIS Evaluation":
    st.header("📋 OASIS Evaluation Formatter")
    uploaded = st.file_uploader("Upload your raw OASIS CSV", type="csv", key="oasis")
    if not uploaded:
        st.stop()

    df = pd.read_csv(uploaded, dtype=str)

    # 自动把 "Course ID"→"course_id", "1 Question Number"→"q1_question_number", …
    def rename_oasis(col: str) -> str:
        col = col.strip()
        m = re.match(r"^(\d+)\s+(.+)$", col)
        if m:
            num, rest = m.groups()
            return f"q{num}_{rest.lower().replace(' ', '_')}"
        return col.lower().replace(" ", "_")

    df.columns = [rename_oasis(c) for c in df.columns]

    # build master_cols
    front = [
        "record_id","course_id","department","course","location",
        "start_date","end_date","course_type","student","student_username",
        "student_external_id","student_designation","student_email",
        "student_aamc_id","student_usmle_id","student_gender","student_level",
        "student_default_classification","evaluator","evaluator_username",
        "evaluator_external_id","evaluator_email","evaluator_gender",
        "who_completed","evaluation","form_record","submit_date"
    ]
    q_sufs = [
        "question_number","question_id","question","answer_text",
        "multiple_choice_order","multiple_choice_value","multiple_choice_label"
    ]
    questions = [f"q{i}_{s}" for i in range(1,24) for s in q_sufs]
    tail = ["oasis_eval_complete"]
    master_cols = front + questions + tail

    # reorder (will KeyError if you missed any)
    df = df.reindex(columns=master_cols)

    # inject REDCap fields
    df["record_id"]                = df["student_external_id"]
    df["redcap_repeat_instrument"] = "oasis_eval"
    df["redcap_repeat_instance"]   = df.groupby("record_id").cumcount() + 1

    # final column order
    keep_front = ["record_id","redcap_repeat_instrument","redcap_repeat_instance"]
    rest       = [c for c in master_cols if c not in keep_front]
    df = df.reindex(columns=keep_front + rest)


    # 8) remove student & location
    df = df.drop(columns=["student","location","start_date","end_date","location"]) #Cannot have these columns in the repeating instrument. 
                      
    st.dataframe(df, height=400)
    st.download_button(
        "📥 Download formatted OASIS CSV",
        df.to_csv(index=False).encode("utf-8"),
        file_name="oasis_eval_formatted.csv",
        mime="text/csv",
    )


elif instrument == "Checklist Entry":
    st.header("🔖 Checklist Entry Merger")
    uploaded = st.file_uploader(
        "Upload exactly two checklist CSVs",
        type="csv",
        accept_multiple_files=True,
        key="clist"
    )
    if not uploaded:
        st.stop()
    if len(uploaded) != 2:
        st.warning("Please upload *exactly* two CSV files here.")
        st.stop()

    # Read + concat
    dfs = [pd.read_csv(f, dtype=str) for f in uploaded]
    df_cl = pd.concat(dfs, ignore_index=True, sort=False)

    # Rename only your 22 columns
    rename_map = {
        "Student name":           "student_name",
        "External ID":            "external_id",
        "Email":                  "email",
        "Start Date":             "start_date",
        "Location":               "location_cl",
        "Checklist":              "checklist",
        "Checklist status":       "checklist_status",
        "Item":                   "item",
        "Item status":            "item_status",
        "Original/Copy":          "originalcopy",
        "Signed By":              "signed_by",
        "Time Signed":            "time_signed",
        "Verified By":            "verified_by",
        "Verification Comments":  "verification_comments",
        "Verified Date":          "verified_date",
        "Time entered":           "time_entered",
        "Date":                   "date",
        "Times observed":         "times_observed",
        "Is proficient":          "is_proficient",
        "Needs Practice":         "needs_practice",
        "Comments":               "comments",
    }
    df_cl = df_cl.rename(columns=rename_map)

    # Select + reorder
    target = list(rename_map.values())
    df_cl = df_cl[target]

    # Move external_id → record_id up front
    df_cl = df_cl.rename(columns={"external_id": "record_id"})
    cols = ["record_id"] + [c for c in df_cl.columns if c != "record_id"]
    df_cl = df_cl[cols]

    # Add REDCap repeater
    df_cl["redcap_repeat_instrument"] = "checklist_entry"
    df_cl["redcap_repeat_instance"] = df_cl.groupby("record_id").cumcount() + 1

    # Final column order
    all_cols = df_cl.columns.tolist()
    all_cols = [c for c in all_cols if c not in ("redcap_repeat_instrument", "redcap_repeat_instance")]
    all_cols += ["redcap_repeat_instrument", "redcap_repeat_instance"]
    df_cl = df_cl[all_cols]

    # Ensure 'time_entered' is datetime
    df_cl["time_entered"] = pd.to_datetime(df_cl["time_entered"], errors="coerce")

    # Group by record_id and compute max and min start_date
    submitted_max = df_cl.groupby("record_id")["time_entered"].max().dt.strftime("%m-%d-%Y")
    submitted_min = df_cl.groupby("record_id")["time_entered"].min().dt.strftime("%m-%d-%Y")

    # Add empty submitted_ce columns to df_cl so they exist for reordering
    df_cl["submitted_ce"] = ""
    df_cl["submitted_ce_min"] = ""

    # Create summary rows (non-repeating)
    df_summary = pd.DataFrame({
        "record_id": submitted_max.index,
        "submitted_ce": submitted_max.values,
        "submitted_ce_min": submitted_min.values,
        "redcap_repeat_instrument": "",
        "redcap_repeat_instance": ""
    })

    # Add missing columns to match df_cl
    for col in df_cl.columns:
        if col not in df_summary.columns:
            df_summary[col] = ""

    # Align column order
    df_summary = df_summary[df_cl.columns]

    # Concatenate checklist entries + summary rows
    df_cl = pd.concat([df_cl, df_summary], ignore_index=True)

    # Move key columns to front
    front_cols = ["record_id", "submitted_ce", "submitted_ce_min"]
    rest_cols = [c for c in df_cl.columns if c not in front_cols]
    df_cl = df_cl[front_cols + rest_cols]

    # Now drop unnecessary columns
    df_cl = df_cl.drop(columns=["email", "date", "start_date"])

    # Show + download
    st.dataframe(df_cl, height=400)
    st.download_button(
        "📥 Download formatted checklist CSV",
        df_cl.to_csv(index=False).encode("utf-8"),
        file_name="checklist_entries.csv",
        mime="text/csv",
    )


# ─── NBME Score ─────────────────────────────────────────────────────────────

elif instrument == "NBME Scores":
    st.header("🔖 NBME")

    # upload exactly one Excel file
    nbme_file = st.file_uploader(
        "Upload exactly one NBME XLSX",
        type=["xlsx"],
        accept_multiple_files=False,
        key="nbme"
    )
    if not nbme_file:
        st.stop()

    # read the specific worksheet
    df_nbme = pd.read_excel(nbme_file, sheet_name="GradeBook", dtype=str)

    # rename only the nine columns you need
    rename_map_nbme = {
        "Student":                        "student_nbme",
        "Email":                          "email_nbme",
        "Username":                       "username",
        "External ID":                    "record_id",
        "Student Level":                  "student_level_nbme",
        "Location":                       "location_nbme",
        "Start Date":                     "start_date_nbme",
        "NBME Exam - Percentage Score":   "nbme",
        "NBME Exam Grade":                "grade_nbme",
        "Final Course Grade":             "final_course_grade",
    }
    df_nbme = df_nbme.rename(columns=rename_map_nbme)

    # keep only those nine columns, in that order
    df_nbme = df_nbme[list(rename_map_nbme.values())]

    # move external_id → record_id up front
    df_nbme = df_nbme.rename(columns={"external_id": "record_id"})
    cols = ["record_id"] + [c for c in df_nbme.columns if c != "record_id"]
    df_nbme = df_nbme[cols]
    
    # add REDCap repeater
    df_nbme["redcap_repeat_instrument"] = "oasis_eval"
    df_nbme["redcap_repeat_instance"]   = df_nbme.groupby("record_id").cumcount() + 1
    
    # preview + download
    st.dataframe(df_nbme, height=400)
    st.download_button(
        "📥 Download formatted NBME XLSX → CSV",
        df_nbme.to_csv(index=False).encode("utf-8"),
        file_name="nbme_scores_formatted.csv",
        mime="text/csv",
    )

elif instrument == "Preceptor Matching":
    st.header("🔖 Preceptor Matching")

    # upload exactly one CSV
    preceptor_file = st.file_uploader(
        "Upload exactly one Preceptor Matching CSV",
        type=["csv"],
        accept_multiple_files=False,
        key="preceptor"
    )
    if not preceptor_file:
        st.stop()

    # read
    df_pmx = pd.read_csv(preceptor_file, dtype=str)

    # drop the unwanted Delete column
    if "Delete" in df_pmx.columns:
        df_pmx = df_pmx.drop(columns=["Delete"])

    # rename only the REDCap-friendly columns
    rename_map = {
        "Start Date":                    "start_date",
        "End Date":                      "end_date",
        "Location":                      "location",
        "Faculty Name":                  "faculty_name",
        "Faculty Username":              "faculty_username",
        "Faculty External ID":           "faculty_external_id",
        "Faculty Email":                 "faculty_email",
        "Type of Association":           "type_of_association",
        "Student Name":                  "student_name",
        "Student Username":              "student_username",
        "Student External ID":           "record_id",
        "Student Email":                 "student_email",
        "Evaluation Period Start Date":  "eval_period_start_date",
        "Evaluation Period End Date":    "eval_period_end_date",
        "Classification":                "classification",
        "Student Activity":              "student_activity",
        "Manual Evaluations":            "manual_evaluations",
    }
    df_pmx = df_pmx.rename(columns=rename_map)

    # keep only those columns, in that exact order
    df_pmx = df_pmx[list(rename_map.values())]

    # move record_id to front
    df_pmx = df_pmx[["record_id"] + [c for c in df_pmx.columns if c != "record_id"]]

    # add REDCap repeater fields
    df_pmx["redcap_repeat_instrument"] = "oasis_eval"
    df_pmx["redcap_repeat_instance"]   = df_pmx.groupby("record_id").cumcount() + 1

    df_pmx = df_pmx.drop(columns=["start_date","end_date","location","student_name","student_username","student_email"])

    # ─── normalize manual_evaluations to one per row ────────────────────
    # split on "|" into lists
    df_pmx["manual_evaluations"] = df_pmx["manual_evaluations"] \
        .fillna("") \
        .str.split("|")
    
    # explode so each list element gets its own row
    df_pmx = df_pmx.explode("manual_evaluations")
    
    # remove leading "*" and any extra whitespace
    df_pmx["manual_evaluations"] = df_pmx["manual_evaluations"] \
        .str.lstrip("*") \
        .str.strip()

        # ─── drop unwanted categories ───────────────────────────────────────
    to_drop = ["Clinical Teaching Eval", "Mid-Cycle Feedback"]
    df_pmx = df_pmx[~df_pmx["manual_evaluations"].isin(to_drop)]

    # get all unique manual_evaluations values
    opts = df_pmx["manual_evaluations"].dropna().unique().tolist()
    
    # multiselect defaulting to all, so you can deselect any you don’t want
    selected = st.multiselect(
        "Filter by manual_evaluations:",
        options=opts,
        default=opts
    )
    
    # filter the DataFrame to only those values
    df_pmx = df_pmx[df_pmx["manual_evaluations"].isin(selected)]


    # preview + download
    st.dataframe(df_pmx, height=400)
    st.download_button(
        "📥 Download formatted Preceptor Matching CSV",
        df_pmx.to_csv(index=False).encode("utf-8"),
        file_name="preceptor_matching_formatted.csv",
        mime="text/csv",
    )

elif instrument == "Email Record Mapper":
    st.header("📧 Email Record Mapper")

    # Upload exactly one Roster CSV
    roster_file = st.file_uploader(
        "Upload a Roster CSV",
        type=["csv"],
        accept_multiple_files=False,
        key="roster_upload"
    )
    if not roster_file:
        st.stop()

    # Read the CSV
    df_roster = pd.read_csv(roster_file, dtype=str)

    # Rename only the needed columns
    rename_map = {
        "Email Address": "email",
        "External ID": "record_id",
    }
    df_roster = df_roster.rename(columns=rename_map)

    # Keep only relevant columns (you can customize this list)
    keep_cols = ["record_id", "email"]
    df_roster = df_roster[keep_cols]

    # Remove rows with missing emails or record_id
    df_roster = df_roster.dropna(subset=["record_id", "email"])

    # Optional: Sort by student_name or record_id
    df_roster = df_roster.sort_values(by="email")

    # Preview and download
    st.dataframe(df_roster, height=400)
    # Create a Word doc in memory
    doc = Document()
    doc.add_heading('REDCap Dropdown: record_id, email', level=1)
    
    # Add each line as plain text
    for _, row in df_roster.iterrows():
        record_id = str(row['record_id']).strip()
        email = str(row['email']).strip()
        doc.add_paragraph(f"{record_id}, {email}")
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    st.download_button(
        label="📥 Download REDCap Dropdown (Word)",
        data=doc_io,
        file_name="email_roster_dropdown.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
elif instrument == "Weekly Quiz Reports":
    st.header("🔖 Weekly Quiz Reports")

    # 1) Upload exactly four CSVs
    uploaded = st.file_uploader(
        "Upload exactly four Weekly Quiz CSVs",
        type=["csv"],
        accept_multiple_files=True,
        key="weekly_quiz"
    )
    if not uploaded:
        st.stop()
    if len(uploaded) != 4:
        st.warning("Please upload *exactly* four CSV files here.")
        st.stop()

    for file in uploaded:
        df = pd.read_csv(file, dtype=str)
    
        # Identify the quiz week from the filename
        week = None
        if 'week 1' in file.name.lower():
            week = 1
        elif 'week 2' in file.name.lower():
            week = 2
        elif 'week 3' in file.name.lower():
            week = 3
        elif 'week 4' in file.name.lower():
            week = 4
        else:
            st.warning(f"Could not identify week from filename: {file.name}")
            continue
    
        # Rename columns
        quiz_score_column = f"quiz{week}"
        quiz_late_column  = f"quiz_{week}_late"
        df.rename(columns={
            "sis_id": "record_id",
            "submitted": quiz_late_column,
            "score": quiz_score_column,
        }, inplace=True)
    
        # Clean record_id
        df["record_id"] = df["record_id"].str.replace(r"@psu\.edu", "", regex=True)
    
        # Convert quiz score to percentage
        df[quiz_score_column] = pd.to_numeric(df[quiz_score_column], errors='coerce')
        df[quiz_score_column] = (df[quiz_score_column] / 20) * 100
    
        # Parse and localize late date
        df[quiz_late_column] = pd.to_datetime(df[quiz_late_column], errors='coerce')
        if df[quiz_late_column].dt.tz is None:
            df[quiz_late_column] = df[quiz_late_column].dt.tz_localize('UTC').dt.tz_convert('US/Eastern')
        else:
            df[quiz_late_column] = df[quiz_late_column].dt.tz_convert('US/Eastern')
        df[quiz_late_column] = (df[quiz_late_column].dt.normalize() + pd.Timedelta(hours=23, minutes=59)).dt.strftime("%m-%d-%Y %H:%M")
    
        # Keep only what's needed
        df = df[["record_id", quiz_score_column, quiz_late_column]]
    
        # Store in variable like df_1, df_2, etc.
        globals()[f"df_{week}"] = df

    # Outer join them all using record_id
    dfs = [globals().get(f"df_{w}") for w in range(1, 5) if f"df_{w}" in globals()]
    df_quiz_combined = dfs[0]
    for df in dfs[1:]:
        df_quiz_combined = pd.merge(df_quiz_combined, df, on="record_id", how="outer")
    
    # Order columns
    final_columns = ["record_id"] + [f"quiz{w}" for w in range(1, 5)] + [f"quiz_{w}_late" for w in range(1, 5)]
    df_quiz_combined = df_quiz_combined.reindex(columns=final_columns)
    
    # Preview + download
    st.dataframe(df_quiz_combined, height=400)
    st.download_button(
        "📥 Download formatted Weekly Quiz CSV",
        df_quiz_combined.to_csv(index=False).encode("utf-8"),
        file_name="weekly_quiz_formatted.csv",
        mime="text/csv",
    )


elif instrument == "SDOH Form":
    st.header("📧 SDOH Form")

    # Upload exactly one CSV
    roster_file = st.file_uploader(
        "Upload a Roster CSV",
        type=["csv"],
        accept_multiple_files=False,
        key="roster_upload"
    )
    if not roster_file:
        st.stop()

    # Read the CSV
    df = pd.read_csv(roster_file, dtype=str)

    # Only keep the two columns you care about
    cols = ["email_2", "social_drivers_of_health_sdoh_assessment_form_timestamp", "social_drivers_of_health_sdoh_assessment_form_complete"]
    missing = set(cols) - set(df.columns)
    if missing:
        st.error(f"Missing expected columns: {', '.join(missing)}")
        st.stop()
    df = df[cols].copy()

    # Convert the SDOH-complete column to numeric, so max() works
    df["social_drivers_of_health_sdoh_assessment_form_complete"] = pd.to_numeric(
        df["social_drivers_of_health_sdoh_assessment_form_complete"],
        errors="coerce"
    )

    # Keep the row with the max value for each email_2
    df = df.sort_values("social_drivers_of_health_sdoh_assessment_form_complete", ascending=False)
    df = df.drop_duplicates(subset=["email_2"], keep="first")

    # Rename columns
    df_grouped = df.rename(columns={
        "email_2": "record_id",
        "social_drivers_of_health_sdoh_assessment_form_complete": "sdohass",
        "social_drivers_of_health_sdoh_assessment_form_timestamp": "submitted_sdoh"
    })

    # Format the timestamp if possible
    try:
        df_grouped["submitted_sdoh"] = pd.to_datetime(df_grouped["submitted_sdoh"]).dt.strftime("%m-%d-%Y")
    except Exception:
        pass  # Skip formatting if parsing fails

    # Preview in Streamlit
    st.dataframe(df_grouped, height=400)

    # Drop rows with missing record_id
    df_grouped = df_grouped[df_grouped["record_id"].notna() & (df_grouped["record_id"].str.strip() != "")]

    
    # Offer as CSV download
    csv_bytes = df_grouped.to_csv(index=False).encode("utf-8")
    st.download_button(
        label="📥 Download email_2 + SDOH (max) CSV",
        data=csv_bytes,
        file_name="email2_sdoh_max.csv",
        mime="text/csv"
    )


elif instrument == "Developmental Assessment Form":
    st.header("📧 Developmental Assessment Form")

    # Upload exactly one CSV
    roster_file = st.file_uploader(
        "Upload a Roster CSV",
        type=["csv"],
        accept_multiple_files=False,
        key="roster_upload"
    )
    if not roster_file:
        st.stop()

    # Read the CSV
    df = pd.read_csv(roster_file, dtype=str)

    # Only keep the columns you care about
    cols = ["email_2", "developmental_assessment_of_patient_timestamp", "developmental_assessment_of_patient_complete"]
    missing = set(cols) - set(df.columns)
    if missing:
        st.error(f"Missing expected columns: {', '.join(missing)}")
        st.stop()
    df = df[cols].copy()

    # Convert the complete column to numeric
    df["developmental_assessment_of_patient_complete"] = pd.to_numeric(
        df["developmental_assessment_of_patient_complete"],
        errors="coerce"
    )

    # Keep the row with the max complete value per email_2
    df = df.sort_values("developmental_assessment_of_patient_complete", ascending=False)
    df = df.drop_duplicates(subset=["email_2"], keep="first")

    # Rename columns
    df_grouped = df.rename(columns={
        "email_2": "record_id",
        "developmental_assessment_of_patient_complete": "devass",
        "developmental_assessment_of_patient_timestamp": "submitted_dev"
    })

    # Format the timestamp column
    try:
        df_grouped["submitted_dev"] = pd.to_datetime(df_grouped["submitted_dev"]).dt.strftime("%m-%d-%Y")
    except Exception:
        pass  # Skip formatting if parsing fails

    # Preview in Streamlit
    st.dataframe(df_grouped, height=400)

    # Drop rows with missing record_id
    df_grouped = df_grouped[df_grouped["record_id"].notna() & (df_grouped["record_id"].str.strip() != "")]

    
    # Offer as CSV download
    csv_bytes = df_grouped.to_csv(index=False).encode("utf-8")
    st.download_button(
        label="📥 Download record_id + Developmental (max) CSV",
        data=csv_bytes,
        file_name="email2_dev_max.csv",
        mime="text/csv"
    )

elif instrument == "Documentation Submission #1":
    st.header("📧 Documentation Submission #1")

    # Upload exactly one CSV
    roster_file = st.file_uploader(
        "Upload a Roster CSV",
        type=["csv"],
        accept_multiple_files=False,
        key="roster_upload"
    )
    if not roster_file:
        st.stop()

    # Read the CSV
    df = pd.read_csv(roster_file, dtype=str)

    # Only keep the columns you care about
    cols = [
        "documentation_submission_1_timestamp", "email_2", "age_v1", "visit_date_v1", "setting_v1", "chief_v1", "cc_v1",
        "historian_v1", "super_clinician_v1", "historyofpresentillness_v1", "reviewofsystems_v1", "hpi_v1", "ros_v1",
        "pmhx_v1", "pshx_v1", "birthhx_v1", "famhx_v1", "socialhx_v1", "meds_v1", "imm_v1", "allg_v1", "diet_v1", "dev_v1",
        "addhx_v1", "soc_hx_features_v1", "all_v1", "med_v1", "temp_v1", "hr_v1", "rr_v1", "pulseox_v1", "sbp_v1", "dbp_v1",
        "weight_v1", "weighttile_v1", "height_v1", "heighttile_v1", "bmi_v1", "bmitile_v1", "vs_v1", "physicalexam_v1",
        "pe_v1", "dxs_v1", "dxstud_v1", "probrep_v1", "probstatement_v1", "mostlikelydiagnosis_v1",
        "seclikelydiagnosis_v1", "thirlikelydiagnosis_v1", "mostlikelydiagnosisj_v1", "seclikelydiagnosisj_v1",
        "thirlikelydiagnosisj_v1", "diffdx_v1", "txplan_v1", "probid_v1", "plan_v1", "grammar_v1", "hpiwordcount_v1",
        "hpiwords_v1", "score_v1", "scorep_v1", "doccomment_v1"
    ]

    missing = set(cols) - set(df.columns)
    if missing:
        st.error(f"Missing expected columns: {', '.join(missing)}")
        st.stop()
    df = df[cols].copy()

    df = df.rename(columns={"email_2": "record_id", "documentation_submission_1_timestamp":"peddoclate1"})

    df["peddoclate1"] = pd.to_datetime(df["peddoclate1"]).dt.strftime("%m-%d-%Y")

    # Preview in Streamlit
    st.dataframe(df, height=400)

    # Drop rows with missing record_id
    df = df[df["record_id"].notna() & (df["record_id"].str.strip() != "")]

    cols = ["record_id"] + [col for col in df.columns if col != "record_id"]
    df = df[cols]

    # Offer as CSV download
    csv_bytes = df.to_csv(index=False).encode("utf-8")
    st.download_button(label="📥 Download Documentation Submission #1",data=csv_bytes,file_name="docsubmit1.csv",mime="text/csv")

elif instrument == "Documentation Submission #2":
    st.header("📧 Documentation Submission #2")

    # Upload exactly one CSV
    roster_file = st.file_uploader("Upload a Roster CSV",type=["csv"],accept_multiple_files=False,key="roster_upload")
    
    if not roster_file:
        st.stop()

    # Read the CSV
    df = pd.read_csv(roster_file, dtype=str)

    # Only keep the columns you care about
    cols = [
    "documentation_submission_2_timestamp", "email_2", "age_v2", "visit_date_v2", "setting_v2", "chief_v2", "cc_v2",
    "historian_v2", "super_clinician_v2", "historyofpresentillness_v2", "reviewofsystems_v2", "hpi_v2", "ros_v2",
    "pmhx_v2", "pshx_v2", "birthhx_v2", "famhx_v2", "socialhx_v2", "meds_v2", "imm_v2", "allg_v2", "diet_v2", "dev_v2",
    "addhx_v2", "soc_hx_features_v2", "all_v2", "med_v2", "temp_v2", "hr_v2", "rr_v2", "pulseox_v2", "sbp_v2", "dbp_v2",
    "weight_v2", "weighttile_v2", "height_v2", "heighttile_v2", "bmi_v2", "bmitile_v2", "vs_v2", "physicalexam_v2",
    "pe_v2", "dxs_v2", "dxstud_v2", "probrep_v2", "probstatement_v2", "mostlikelydiagnosis_v2",
    "seclikelydiagnosis_v2", "thirlikelydiagnosis_v2", "mostlikelydiagnosisj_v2", "seclikelydiagnosisj_v2",
    "thirlikelydiagnosisj_v2", "diffdx_v2", "txplan_v2", "probid_v2", "plan_v2", "grammar_v2", "hpiwordcount_v2",
    "hpiwords_v2", "score_v2", "scorep_v2", "doccomment_v2"]

    missing = set(cols) - set(df.columns)
    if missing:
        st.error(f"Missing expected columns: {', '.join(missing)}")
        st.stop()
    df = df[cols].copy()

    df = df.rename(columns={"email_2": "record_id", "documentation_submission_2_timestamp":"peddoclate2"})

    df["peddoclate2"] = pd.to_datetime(df["peddoclate2"]).dt.strftime("%m-%d-%Y")

    # Preview in Streamlit
    st.dataframe(df, height=400)

    # Drop rows with missing record_id
    df = df[df["record_id"].notna() & (df["record_id"].str.strip() != "")]

    cols = ["record_id"] + [col for col in df.columns if col != "record_id"]
    df = df[cols]

    # Offer as CSV download
    csv_bytes = df.to_csv(index=False).encode("utf-8")
    st.download_button(label="📥 Download Documentation Submission #2",data=csv_bytes,file_name="docsubmit1.csv",mime="text/csv")
    
elif instrument == "Roster_HMC":
    st.header("🔖 Roster_HMC")

    # upload exactly one CSV
    roster_file = st.file_uploader(
        "Upload exactly one Roster CSV",
        type=["csv"],
        accept_multiple_files=False,
        key="roster"
    )
    if not roster_file:
        st.stop()

    # read as CSV
    df_roster = pd.read_csv(roster_file, dtype=str)

    df_roster.columns = df_roster.columns.str.strip()

    # map your columns to REDCap-friendly names
    rename_map = {
        "#":                              "row_number",
        "Student":                        "student",
        "Legal Name":                     "legal_name",
        "Previous Name":                  "previous_name",
        "Username":                       "username",
        "Confidential":                   "confidential",
        "External ID":                    "record_id",
        "Email Address":                  "email",
        "Phone":                          "phone",
        "Pager":                          "pager",
        "Mobile":                         "mobile",
        "Gender":                         "gender",
        "Pronouns":                       "pronouns",
        "Ethnicity":                      "ethnicity",
        "Designation":                    "designation",
        "AAMC ID":                        "aamc_id",
        "USMLE ID":                       "usmle_id",
        "Home School":                    "home_school",
        "Campus":                         "campus",
        "Date of Birth":                  "date_of_birth",
        "Emergency Contact":              "emergency_contact",
        "Emergency Phone":                "emergency_phone",
        "Primary Academic Department":    "primary_academic_department",
        "Secondary Academic Department":  "secondary_academic_department",
        "Academic Type":                  "academic_type",
        "Primary Site":                   "primary_site",
        "NBME":                           "nbme_score",
        "PSU ID":                         "psu_id",
        "Productivity Specialty":         "productivity_specialty",
        "Grade":                          "grade",
        "Status":                         "status",
        "Student Level":                  "student_level",
        "Track":                          "track",
        "Location":                       "location",
        "Start Date":                     "start_date",
        "End Date":                       "end_date",
        "Weeks":                          "weeks",
        "Credits":                        "credits",
        "Enrolled":                       "enrolled",
        "Actions":                        "actions",
        "Aprv By":                     "approved_by"
    }
    df_roster = df_roster.rename(columns=rename_map)

    # keep only those renamed columns (in this exact order)
    df_roster = df_roster[list(rename_map.values())]

    # move record_id to the front
    cols = ["record_id"] + [c for c in df_roster.columns if c != "record_id"]
    df_roster = df_roster[cols]

    # add REDCap repeater - dont need
    #df_roster["redcap_repeat_instrument"] = "roster"
    #df_roster["redcap_repeat_instance"]   = df_roster.groupby("record_id").cumcount() + 1

    # ─── split “student” into last_name / first_name ─────────────────────────
    # 1) drop everything after the semicolon
    name_only = df_roster["student"].str.split(";", n=1).str[0]
    # 2) split on comma into last / first
    parts = name_only.str.split(",", n=1, expand=True)
    df_roster["lastname"]  = parts[0].str.strip()
    df_roster["firstname"] = parts[1].str.strip()

    df_roster["name"] = df_roster["firstname"] + " " + df_roster["lastname"]

    df_roster["legal_name"] = df_roster["lastname"] + ", " + df_roster["firstname"] + " (MD)" 

    df_roster["email_2"] = df_roster["record_id"] + "@psu.edu"

    #legal name ... legal_name
    
    # 3) (optional) drop the original combined column
    renamed_cols_a = ["row_number","student","previous_name","username","confidential","phone","pager","mobile","gender","pronouns","ethnicity","designation","aamc_id","usmle_id","home_school"]
    renamed_cols_b = ["campus","date_of_birth","emergency_contact","emergency_phone","primary_academic_department","secondary_academic_department","academic_type","primary_site","nbme_score"]
    renamed_cols_c = ["productivity_specialty","grade","status","student_level","weeks","credits","enrolled","actions","approved_by"]

    renamed_cols = renamed_cols_a + renamed_cols_b + renamed_cols_c
    
    # 0) ensure start_date is a true datetime
    df_roster["start_date"] = pd.to_datetime(df_roster["start_date"], errors='coerce')
    
    # 1) grab each unique date, sorted oldest → newest
    unique_dates = sorted(df_roster["start_date"].dropna().unique())
    
    # 2) for each one, make a new column rot_date_#
    for idx, dt in enumerate(unique_dates, 1):
        df_roster[f"rot_date_{idx}"] = df_roster["start_date"].apply(lambda x: dt.strftime("%m-%d-%Y") if pd.notna(x) and x == dt else "")

    # 3) build a mapping from date → rotation code
    rotation_map = {dt: f"r{idx}" for idx, dt in enumerate(unique_dates, 1)}
    
    # 4) assign each student’s rotation1 based on their start_date
    df_roster["rotation1"] = df_roster["start_date"].map(rotation_map)

    df_roster["rotation"] = df_roster["start_date"].map(rotation_map)

    # 3) now drop your old columns
    df_roster.drop(columns=renamed_cols, errors="ignore", inplace=True)

    #DUE DATES
    
    # ─── 1) Ensure start_date and end_date are datetime ─────────────────────────
    df_roster["start_date"] = pd.to_datetime(df_roster["start_date"], infer_datetime_format=True)
    df_roster["end_date"]   = pd.to_datetime(df_roster["end_date"], infer_datetime_format=True)
    
    # ─── 2) Compute first Sunday on/after start_date ────────────────────────────
    days_to_sunday = (6 - df_roster["start_date"].dt.weekday) % 7
    first_sunday   = df_roster["start_date"] + pd.to_timedelta(days_to_sunday, unit="D")
    
    # ─── 3) Create quiz_due_1 … quiz_due_4 ──────────────────────────────────────
    for n in range(1, 5):
        df_roster[f"quiz_due_{n}"] = first_sunday + pd.Timedelta(weeks=(n - 1))
    
    # ─── 4) Alias assignment & doc-assignment due dates ─────────────────────────
    df_roster["ass_middue_date"]   = df_roster["quiz_due_2"]
    df_roster["ass_due_date"]      = df_roster["quiz_due_4"]
    df_roster["docass_due_date_1"] = df_roster["quiz_due_2"]
    df_roster["docass_due_date_2"] = df_roster["quiz_due_4"]
    
    # ─── 5) Grade due date: 6 weeks after end_date ──────────────────────────────
    df_roster["grade_due_date"] = df_roster["end_date"] + pd.Timedelta(weeks=6)

    # ─── 6) Normalize all due dates to 23:59 with no seconds ─────────────────
    due_cols = [
        "quiz_due_1","quiz_due_2","quiz_due_3","quiz_due_4",
        "ass_middue_date","ass_due_date",
        "docass_due_date_1","docass_due_date_2",
        "grade_due_date"
    ]
    
    for col in due_cols:
        #df_roster[col] = (df_roster[col].dt.normalize() + pd.Timedelta(hours=23, minutes=59)).dt.strftime("%Y-%m-%d %H:%M")
        #df_roster[col] = (df_roster[col].dt.normalize() + pd.Timedelta(hours=23, minutes=59)).dt.strftime("%m/%d/%Y 23:59")
        #df_roster[col] = (df_roster[col].dt.normalize() + pd.Timedelta(hours=23, minutes=59)).dt.strftime("%-m/%-d/%Y 23:59")
        df_roster[col] = (df_roster[col].dt.normalize() + pd.Timedelta(hours=23, minutes=59)).dt.strftime("%m-%d-%Y 23:59")

    df_roster["start_date"] = df_roster["start_date"].dt.strftime("%m-%d-%Y")
    df_roster["end_date"] = df_roster["end_date"].dt.strftime("%m-%d-%Y")
    
    # preview + download
    st.dataframe(df_roster, height=400)
    
    st.download_button("📥 Download formatted Roster CSV",df_roster.to_csv(index=False).encode("utf-8"),file_name="roster_formatted.csv",mime="text/csv")

elif instrument == "Roster_KP":
    st.header("🔖 Roster KP")

    # upload exactly one CSV
    roster_file = st.file_uploader("Upload exactly one Roster CSV",type=["csv"],accept_multiple_files=False,key="roster")
    if not roster_file:
        st.stop()

    # read as CSV
    df_roster = pd.read_csv(roster_file, dtype=str)

    df_roster.columns = df_roster.columns.str.strip()

    # map your columns to REDCap-friendly names
    rename_map = {
        "#":                              "row_number",
        "Student":                        "student",
        "Legal Name":                     "legal_name",
        "Previous Name":                  "previous_name",
        "Username":                       "username",
        "Confidential":                   "confidential",
        "External ID":                    "record_id",
        "Email Address":                  "email",
        "Phone":                          "phone",
        "Pager":                          "pager",
        "Mobile":                         "mobile",
        "Gender":                         "gender",
        "Pronouns":                       "pronouns",
        "Ethnicity":                      "ethnicity",
        "Designation":                    "designation",
        "AAMC ID":                        "aamc_id",
        "USMLE ID":                       "usmle_id",
        "Home School":                    "home_school",
        "Campus":                         "campus",
        "Date of Birth":                  "date_of_birth",
        "Emergency Contact":              "emergency_contact",
        "Emergency Phone":                "emergency_phone",
        "Primary Academic Department":    "primary_academic_department",
        "Secondary Academic Department":  "secondary_academic_department",
        "Academic Type":                  "academic_type",
        "Primary Site":                   "primary_site",
        "NBME":                           "nbme_score",
        "PSU ID":                         "psu_id",
        "Productivity Specialty":         "productivity_specialty",
        "Grade":                          "grade",
        "Status":                         "status",
        "Student Level":                  "student_level",
        "Track":                          "track",
        "Location":                       "location",
        "Start Date":                     "start_date",
        "End Date":                       "end_date",
        "Weeks":                          "weeks",
        "Credits":                        "credits",
        "Enrolled":                       "enrolled",
        "Actions":                        "actions",
        "Aprv By":                     "approved_by"
    }
    df_roster = df_roster.rename(columns=rename_map)

    # keep only those renamed columns (in this exact order)
    df_roster = df_roster[list(rename_map.values())]

    # move record_id to the front
    cols = ["record_id"] + [c for c in df_roster.columns if c != "record_id"]
    df_roster = df_roster[cols]

    # ─── split “student” into last_name / first_name ─────────────────────────
    # 1) drop everything after the semicolon
    name_only = df_roster["student"].str.split(";", n=1).str[0]
    # 2) split on comma into last / first
    parts = name_only.str.split(",", n=1, expand=True)
    df_roster["lastname"]  = parts[0].str.strip()
    df_roster["firstname"] = parts[1].str.strip()

    df_roster["name"] = df_roster["firstname"] + " " + df_roster["lastname"]

    df_roster["legal_name"] = df_roster["lastname"] + ", " + df_roster["firstname"] + " (MD)" 

    df_roster["email_2"] = df_roster["record_id"] + "@psu.edu"

    #legal name ... legal_name
    
    # 3) (optional) drop the original combined column
    renamed_cols_a = ["row_number","student","previous_name","username","confidential","phone","pager","mobile","gender","pronouns","ethnicity","designation","aamc_id","usmle_id","home_school"]
    renamed_cols_b = ["campus","date_of_birth","emergency_contact","emergency_phone","primary_academic_department","secondary_academic_department","academic_type","primary_site","nbme_score"]
    renamed_cols_c = ["productivity_specialty","grade","status","student_level","weeks","credits","enrolled","actions","approved_by"]

    renamed_cols = renamed_cols_a + renamed_cols_b + renamed_cols_c
    
    # 0) ensure start_date is a true datetime
    df_roster["start_date"] = pd.to_datetime(df_roster["start_date"], errors='coerce')
    
    # 1) grab each unique date, sorted oldest → newest
    unique_dates = sorted(df_roster["start_date"].dropna().unique())
    
    # 2) for each one, make a new column rot_date_#
    for idx, dt in enumerate(unique_dates, 1):
        df_roster[f"rot_date_{idx}"] = df_roster["start_date"].apply(lambda x: dt.strftime("%m-%d-%Y") if pd.notna(x) and x == dt else "")

    # 3) build a mapping from date → rotation code
    rotation_map = {dt: f"r{idx}" for idx, dt in enumerate(unique_dates, 1)}
    
    # 4) assign each student’s rotation1 based on their start_date
    df_roster["rotation1"] = "KPLIC"

    df_roster["rotation"] = "KPLIC"

    # 3) now drop your old columns
    df_roster.drop(columns=renamed_cols, errors="ignore", inplace=True)

    #DUE DATES
    
    # ─── 1) Ensure start_date and end_date are datetime ─────────────────────────
    df_roster["start_date"] = pd.to_datetime(df_roster["start_date"], infer_datetime_format=True)
    df_roster["end_date"]   = pd.to_datetime(df_roster["end_date"], infer_datetime_format=True)
    
    # ─── 2) Compute first Sunday on/after start_date ────────────────────────────
    days_to_sunday = (6 - df_roster["start_date"].dt.weekday) % 7
    first_sunday   = df_roster["start_date"] + pd.to_timedelta(days_to_sunday, unit="D")
    
    # ─── 3) Create quiz_due_1 … quiz_due_4 ──────────────────────────────────────
    for n in range(1, 5):
        df_roster[f"quiz_due_{n}"] = first_sunday + pd.Timedelta(weeks=(n - 1))
    
    # ─── 4) Alias assignment & doc-assignment due dates ─────────────────────────
    df_roster["ass_middue_date"]   = df_roster["quiz_due_2"]
    df_roster["ass_due_date"]      = df_roster["quiz_due_4"]
    df_roster["docass_due_date_1"] = df_roster["quiz_due_2"]
    df_roster["docass_due_date_2"] = df_roster["quiz_due_4"]
    
    # ─── 5) Grade due date: 6 weeks after end_date ──────────────────────────────
    df_roster["grade_due_date"] = df_roster["end_date"] + pd.Timedelta(weeks=6)

    # ─── 6) Normalize all due dates to 23:59 with no seconds ─────────────────
    due_cols = [
        "quiz_due_1","quiz_due_2","quiz_due_3","quiz_due_4",
        "ass_middue_date","ass_due_date",
        "docass_due_date_1","docass_due_date_2",
        "grade_due_date"
    ]
    
    for col in due_cols:
        df_roster[col] = (df_roster[col].dt.normalize() + pd.Timedelta(hours=23, minutes=59)).dt.strftime("%m-%d-%Y 23:59")

    df_roster["start_date"] = df_roster["start_date"].dt.strftime("%m-%d-%Y")
    df_roster["end_date"] = df_roster["end_date"].dt.strftime("%m-%d-%Y")
    
    # preview + download
    st.dataframe(df_roster, height=400)
    
    st.download_button("📥 Download formatted Roster CSV",df_roster.to_csv(index=False).encode("utf-8"),file_name="roster_formatted.csv",mime="text/csv")


